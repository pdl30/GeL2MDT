"""Copyright (c) 2018 Great Ormond Street Hospital for Children NHS Foundation
Trust & Birmingham Women's and Children's NHS Foundation Trust

Permission is hereby granted, free of charge, to any person obtaining a copy of
this software and associated documentation files (the "Software"), to deal in
the Software without restriction, including without limitation the rights to
use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies
of the Software, and to permit persons to whom the Software is furnished to do
so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""
from .models import *
import csv
import xlsxwriter
import io
from docx import Document, oxml, opc
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE
from django.conf import settings
import os
from docx.shared import Pt, Inches, RGBColor, Cm
from datetime import datetime


def write_mdt_export(mdt_instance, mdt_reports):
    '''
    Writes a summary of the cases which are being brought to MDT
    :param writer: CSV file writer
    :param mdt_instance:  MDT instance
    :param mdt_reports: List of reports which are present in MDT
    :return: CSV file Writer
    '''
    failed_reports = []
    for report in mdt_reports:
        proband_variants = ProbandVariant.objects.filter(interpretation_report=report.interpretation_report)
        for proband_variant in proband_variants:
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant is None:
                failed_reports.append(report.interpretation_report.ir_family.ir_family_id)

    if failed_reports:
        failed_reports_formatted = ' '.join(list(set(failed_reports)))
        raise ValueError(f"Transcripts have not been selected for the following reports: {failed_reports_formatted}")

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet()
    # set formatting
    header_format = workbook.add_format({'bold': 1})
    vcenter_format = workbook.add_format({'valign': 'vcenter'})
    vcenter_date_format = workbook.add_format({'valign': 'vcenter', 'num_format': 'mm/dd/yyyy'})
    # write headings
    if mdt_instance.sample_type == 'raredisease':
        worksheet.write('A1', 'GEL ID', header_format)
        worksheet.write('B1', 'CIP ID', header_format)
        worksheet.write('C1', 'GMC', header_format)
        worksheet.write('D1', 'Forename', header_format)
        worksheet.write('E1', 'Surname', header_format)
        worksheet.write('F1', 'Sex', header_format)
        worksheet.write('G1', 'DOB', header_format)
        worksheet.write('H1', 'NHS number', header_format)
        worksheet.write('I1', 'Family ID', header_format)
        worksheet.write('J1', 'Clinician', header_format)
        worksheet.write('K1', 'Panel(s)', header_format)
        worksheet.write('L1', 'Variant', header_format)
        worksheet.write('M1', 'Inheritance', header_format)
        worksheet.write('N1', 'Proband zygosity', header_format)
        worksheet.write('O1', 'Maternal zygosity', header_format)
        worksheet.write('P1', 'Paternal zygosity', header_format)
        worksheet.write('Q1', 'Phenotypic fit', header_format)
        worksheet.write('R1', 'Discussion required', header_format)
        worksheet.write('S1', 'Comments', header_format)
        worksheet.set_column('A:A', 10)
        worksheet.set_column('B:B', 10)
        worksheet.set_column('C:C', 10)
        worksheet.set_column('D:F', 10)
        worksheet.set_column('G:G', 10)
        worksheet.set_column('H:H', 10)
        worksheet.set_column('I:I', 10)
        worksheet.set_column('J:J', 10)
        worksheet.set_column('K:L', 40)
        worksheet.set_column('M:M', 10)
        worksheet.set_column('N:N', 18)
        worksheet.set_column('O:O', 18)
        worksheet.set_column('P:P', 30)
        worksheet.set_column('Q:Q', 30)
        worksheet.set_column('R:R', 30)
        worksheet.set_column('S:S', 30)
    elif mdt_instance.sample_type == 'cancer':
        worksheet.write('A1', 'GEL ID', header_format)
        worksheet.write('B1', 'CIP ID', header_format)
        worksheet.write('C1', 'LDP', header_format)
        worksheet.write('D1', 'Forename', header_format)
        worksheet.write('E1', 'Surname', header_format)
        worksheet.write('F1', 'Sex', header_format)
        worksheet.write('G1', 'DOB', header_format)
        worksheet.write('H1', 'NHS number', header_format)
        worksheet.write('I1', 'Family ID', header_format)
        worksheet.write('J1', 'Clinician', header_format)
        worksheet.write('K1', 'Recruiting Disease', header_format)
        worksheet.write('L1', 'Disease subtype', header_format)
        worksheet.set_column('A:A', 10)
        worksheet.set_column('B:B', 10)
        worksheet.set_column('C:C', 10)
        worksheet.set_column('D:F', 10)
        worksheet.set_column('G:G', 10)
        worksheet.set_column('H:H', 10)
        worksheet.set_column('I:I', 10)
        worksheet.set_column('J:J', 10)
        worksheet.set_column('K:L', 10)

    row_count = 2
    if mdt_instance.sample_type == 'cancer':
        for report in mdt_reports:
            worksheet.write('A' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.gel_id,
                            vcenter_format)
            worksheet.write('B' + str(row_count),
                            report.interpretation_report.ir_family.ir_family_id, vcenter_format)
            worksheet.write('C' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.gmc, vcenter_format)
            worksheet.write('D' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.forename,
                            vcenter_format)
            worksheet.write('E' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.surname,
                            vcenter_format)
            worksheet.write('F' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.sex, vcenter_format)
            worksheet.write('G' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(),
                            vcenter_date_format)
            worksheet.write('H' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.nhs_number,
                            vcenter_format)
            worksheet.write('I' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.gel_family_id, vcenter_format)
            worksheet.write('J' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.clinician.name,
                            vcenter_format)
            worksheet.write('K' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.recruiting_disease,
                            vcenter_format)
            worksheet.write('L' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.disease_subtype,
                            vcenter_format)
            row_count += 1
    elif mdt_instance.sample_type == 'raredisease':
        for report in mdt_reports:
            proband_variants = ProbandVariant.objects.filter(interpretation_report=report.interpretation_report)
            panels = InterpretationReportFamilyPanel.objects.filter(ir_family=report.interpretation_report.ir_family)
            pv_output = []
            if proband_variants.exists():
                for proband_variant in proband_variants:
                    transcript = proband_variant.get_transcript()
                    transcript_variant = proband_variant.get_transcript_variant()
                    if transcript and transcript_variant:
                        hgvs_c = None
                        hgvs_p = None
                        hgvs_c_split = transcript_variant.hgvs_c.split(':')
                        hgvs_p_split = transcript_variant.hgvs_p.split(':')
                        if len(hgvs_c_split) > 1:
                            hgvs_c = hgvs_c_split[1]
                        if len(hgvs_p_split) > 1:
                            hgvs_p = hgvs_p_split[1].replace('%3D', '=')
                        pv_output.append({'variant': f'{transcript.gene}, '
                                        f'{hgvs_c}, '
                                        f'{hgvs_p}, ',
                                        'inheritance': f'{proband_variant.inheritance}',
                                        'proband_zygosity': f'{proband_variant.zygosity}',
                                        'mat_zygosity': f'{proband_variant.maternal_zygosity}',
                                        'pat_zygosity': f'{proband_variant.paternal_zygosity}',})
                panel_names = []
                for panel in panels:
                    panel_names.append(f'{panel.panel.panel.panel_name}_'
                                    f'{panel.panel.version_number}')

                v_rows = row_count
                for variant in pv_output:
                    worksheet.write('L' + str(v_rows), variant['variant'])
                    worksheet.write('M' + str(v_rows), variant['inheritance'])
                    worksheet.write('N' + str(v_rows), variant['proband_zygosity'])
                    worksheet.write('O' + str(v_rows), variant['mat_zygosity'])
                    worksheet.write('P' + str(v_rows), variant['pat_zygosity'])
                    worksheet.data_validation('Q' + str(v_rows),
                        {'validate': 'list', 'source': ['Yes', 'No', 'Maybe']})
                    worksheet.data_validation('R' + str(v_rows),
                        {'validate': 'list', 'source': ['Yes', 'No']})
                    v_rows += 1

                if row_count == v_rows - 1:
                    worksheet.write('A' + str(row_count),
                                    report.interpretation_report.ir_family.participant_family.proband.gel_id, vcenter_format)
                    worksheet.write('B' + str(row_count),
                                    report.interpretation_report.ir_family.ir_family_id, vcenter_format)
                    worksheet.write('C' + str(row_count),
                                    report.interpretation_report.ir_family.participant_family.proband.gmc, vcenter_format)
                    worksheet.write('D' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.proband.forename, vcenter_format)
                    worksheet.write('E' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.proband.surname, vcenter_format)
                    worksheet.write('F' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.proband.sex, vcenter_format)
                    worksheet.write('G' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(), vcenter_date_format)
                    worksheet.write('H' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.proband.nhs_number, vcenter_format)
                    worksheet.write('I' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.gel_family_id, vcenter_format)
                    worksheet.write('J' + str(row_count),
                        report.interpretation_report.ir_family.participant_family.clinician.name, vcenter_format)
                    worksheet.write('K' + str(row_count),
                        '\n'.join(panel_names), vcenter_format)
                else:
                    worksheet.merge_range('A' + str(row_count) + ':A' + str(v_rows - 1),
                                        report.interpretation_report.ir_family.participant_family.proband.gel_id,
                                        vcenter_format)
                    worksheet.merge_range('B' + str(row_count) + ':B' + str(v_rows - 1),
                                        report.interpretation_report.ir_family.ir_family_id,
                                        vcenter_format)
                    worksheet.merge_range('C' + str(row_count) + ':C' + str(v_rows - 1),
                                        report.interpretation_report.ir_family.participant_family.proband.gmc,
                                        vcenter_format)
                    worksheet.merge_range('D' + str(row_count) + ':D' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.proband.forename,
                        vcenter_format)
                    worksheet.merge_range('E' + str(row_count) + ':E' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.proband.surname,
                        vcenter_format)
                    worksheet.merge_range('F' + str(row_count) + ':F' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.proband.sex,
                        vcenter_format)
                    worksheet.merge_range('G' + str(row_count) + ':G' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(),
                        vcenter_date_format)
                    worksheet.merge_range('H' + str(row_count) + ':H' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.proband.nhs_number,
                        vcenter_format)
                    worksheet.merge_range('I' + str(row_count) + ':I' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.gel_family_id,
                        vcenter_format)
                    worksheet.merge_range('J' + str(row_count) + ':J' + str(v_rows-1),
                        report.interpretation_report.ir_family.participant_family.clinician.name,
                        vcenter_format)
                    worksheet.merge_range('K' + str(row_count) + ':K' + str(v_rows-1),
                        '\n'.join(panel_names),
                        vcenter_format)

                row_count = v_rows

    workbook.close()
    # rewind the buffer
    output.seek(0)
    return output


def monthly_not_completed():
    all_mdts = MDT.objects.all()
    workbook = xlsxwriter.Workbook("monthly_results.xlsx")
    worksheet = workbook.add_worksheet('Summary')
    months = {1: 'Jan', 2: 'Feb', 3: 'Mar', 4: 'Apr', 5: 'May', 6: 'June', 7: 'July', 8: 'Aug', 9: 'Sep', 10: 'Oct',
              11: 'Nov', 12: 'Dec'}
    years = ['2017', '2018', '2019', '2020']
    month_count = 0
    for year in years:
        for month in months:
            completed_cases = []
            notcompleted_cases = []
            worksheet.write(0, month_count, f"{year}_{months[month]}")
            month_mdts = all_mdts.filter(date_of_mdt__year=year, date_of_mdt__month=month)
            for mdt in month_mdts:
                for report in mdt.mdtreport_set.all():
                    if report.interpretation_report.case_status != 'C':
                        notcompleted_cases.append(report)
                    else:
                        completed_cases.append(report)
            worksheet.write(1, month_count, f"Completed Count: ")
            worksheet.write(2, month_count, len(completed_cases))
            worksheet.write(1, month_count + 1, f"Not Completed Count: ")
            worksheet.write(2, month_count + 1, len(notcompleted_cases))
            if notcompleted_cases:
                worksheet.write(4, month_count + 1, 'Participant IDs')
            row = 5
            for case in notcompleted_cases:
                try:
                    worksheet.write(row, month_count + 1,
                                    f"{case.interpretation_report.ir_family.participant_family.proband.gel_id}; "
                                    f"{case.interpretation_report.ir_family.participant_family.clinician.name}")
                    row += 1
                except Proband.DoesNotExist:
                    pass
            month_count += 2
    workbook.close()
    return workbook


def write_mdt_outcome_template(report):
    """
    :param pk: GEL Interpretationreport instance
    :return: Writes a docx template file for summarising proband MDT outcomes
    """
    #document = Document()
    # footers template, page number setup
    mdm_template_file = os.path.join(os.getcwd(), "gel2mdt/exports_templates/{filename}".format(filename='mdm_outcome_template.docx'))
    document = Document(docx=mdm_template_file)

    document.add_picture(os.path.join(settings.STATIC_DIR, 'nhs_image.png'), height=Inches(1.0)) # image scaled on width axis
    header_image = document.paragraphs[-1]
    header_image.alignment = WD_ALIGN_PARAGRAPH.RIGHT 

    style = document.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)
 
    document.add_heading('Genomics MDM record', 0)

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = table.rows[0].cells[0].paragraphs[0].add_run('THIS IS NOT A DIAGNOSTIC REPORT. UNVALIDATED FINDINGS SHOULD NOT BE USED TO INFORM CLINICAL MANAGEMENT DECISIONS.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'This is a record of unvalidated variants identified through the 100,000 genome project.\n'
        'Class 3 variants are of uncertain clinical significance, future review and diagnostic confirmation may '
        'be appropriate if further evidence becomes available.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)

    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.3)
    paragraph = document.add_paragraph()
    paragraph.add_run()

    table = document.add_table(rows=2, cols=4, style='Table Grid')
    heading_cells = table.rows[0].cells
    heading_cells[0].paragraphs[0].add_run('Patient Name').bold=True
    heading_cells[1].paragraphs[0].add_run('DOB').bold=True
    heading_cells[2].paragraphs[0].add_run('NHS number').bold=True
    heading_cells[3].paragraphs[0].add_run('Local ID').bold=True

    row = table.rows[1].cells
    row[0].text = str(report.ir_family.participant_family.proband.forename) \
                  + ' ' \
                  + str(report.ir_family.participant_family.proband.surname)
    row[1].text = str(report.ir_family.participant_family.proband.date_of_birth.date())
    try:
        row[2].text = report.ir_family.participant_family.proband.nhs_number
    except TypeError:
        row[2].text = ''
    if report.ir_family.participant_family.proband.local_id:
        row[3].text = report.ir_family.participant_family.proband.local_id

    paragraph = document.add_paragraph()
    paragraph.add_run()
    paragraph.add_run('Referring Clinician: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.clinician))
    paragraph.add_run('Department/Hospital: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.proband.gmc))
    paragraph.add_run('Study: ').bold=True
    paragraph.add_run('100,000 genomes (whole genome sequencing)\n')
    paragraph.add_run('OPA ID: ').bold = True
    paragraph.add_run('{}\n'.format(report.ir_family.ir_family_id))
    paragraph.add_run('Family ID: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.gel_family_id))
    paragraph.add_run('Genome Build: ').bold=True
    paragraph.add_run('{}\n\n'.format(report.assembly))

    # paragraph.add_run('Phenotype summary: ').bold=True
    # if sample_info.hpo_terms:
    #    paragraph.add_run('{}\n'.format(', '.join(list(json.loads(sample_info.hpo_terms)))))

    proband_variants = list(ProbandVariant.objects.filter(interpretation_report=report))

    run = paragraph.add_run('MDT:\n')
    run.font.size = Pt(16)
    run.underline = True
    run.bold = True

    if proband_variants:
        run = paragraph.add_run('Variant Outcome Summary:\n')
        run.font.size = Pt(13)
        run.underline = True
        run.bold = True

        table = document.add_table(rows=1, cols=7, style='Table Grid')
        heading_cells = table.rows[0].cells
        run = heading_cells[0].paragraphs[0].add_run('Gene')
        run.bold=True
        run.font.size = Pt(9)
        run = heading_cells[1].paragraphs[0].add_run('HGVSg')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[2].paragraphs[0].add_run('HGVSc')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[3].paragraphs[0].add_run('HGVSp')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[4].paragraphs[0].add_run('Zygosity')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[5].paragraphs[0].add_run('Phenotype Contribution')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[6].paragraphs[0].add_run('Class')
        run.bold = True
        run.font.size = Pt(9)

    for proband_variant in proband_variants:
        cells = table.add_row().cells
        transcript = proband_variant.get_transcript()
        transcript_variant = proband_variant.get_transcript_variant()
        if transcript is None or transcript_variant is None:
            raise ValueError(f"Please select transcripts for all variants before exporting\n")
        rdr = proband_variant.create_rare_disease_report()
        run = cells[0].paragraphs[0].add_run(str(transcript.gene))
        run.font.size = Pt(7)
        run = cells[1].paragraphs[0].add_run(str(transcript_variant.hgvs_g))
        run.font.size = Pt(7)
        run = cells[2].paragraphs[0].add_run(str(transcript_variant.hgvs_c))
        run.font.size = Pt(7)
        run = cells[3].paragraphs[0].add_run(str(transcript_variant.hgvs_p))
        run.font.size = Pt(7)
        run = cells[4].paragraphs[0].add_run(str(proband_variant.zygosity))
        run.font.size = Pt(7)
        run = cells[5].paragraphs[0].add_run(str(rdr.get_contribution_to_phenotype_display()))
        run.font.size = Pt(7)
        run = cells[6].paragraphs[0].add_run(str(rdr.classification))
        run.font.size = Pt(7)

    mdt_linkage_list = MDTReport.objects.filter(interpretation_report=report).values('MDT')
    mdt = MDT.objects.filter(id__in=mdt_linkage_list).order_by('-date_of_mdt').first()

    paragraph = document.add_paragraph()

    paragraph.add_run('MDT Date: ').bold = True
    paragraph.add_run('{}\n'.format(mdt.date_of_mdt.date()))
    paragraph.add_run('MDT Attendees: ').bold = True
    clinicians = Clinician.objects.filter(mdt=mdt.id).values_list('name', flat=True)
    clinical_scientists = ClinicalScientist.objects.filter(mdt=mdt.id).values_list('name', flat=True)
    other_staff = OtherStaff.objects.filter(mdt=mdt.id).values_list('name', flat=True)

    attendees = list(clinicians) + list(clinical_scientists) + list(other_staff)
    paragraph.add_run('{}\n\n'.format(', '.join(attendees)))
    paragraph.add_run()
    run = paragraph.add_run('Discussion:\n')
    run.font.size = Pt(13)
    run.underline = True
    run.bold = True
    paragraph.add_run('{}\n\n'.format(report.ir_family.participant_family.proband.discussion.rstrip()))
    run = paragraph.add_run('Action:\n')
    run.font.size = Pt(13)
    run.underline = True
    run.bold = True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.proband.action.rstrip()))
    return document, mdt


def write_gtab_template(report):
    '''
    Given a Cancer report, write a report template for GTAB
    :param report: GELInterpretation instance
    :return: docx document to be exported
    '''
    proband_variants = list(ProbandVariant.objects.filter(interpretation_report=report))
    
    # footers template, page number setup
    template_file = os.path.join(os.getcwd(), "gel2mdt/exports_templates/{filename}".format(filename='gtab_template.docx'))
    document = Document(docx=template_file) 

    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    #document.add_picture(os.path.join(settings.STATIC_DIR, 'nhs_image.png'), height=Inches(0.63), width=Inches(2.39))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Page 1.
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    heading_cells = table.rows[0].cells
    run = heading_cells[0].paragraphs[0].add_run(
        'GENOMICS TUMOUR ADVISORY BOARD (GTAB) SUMMARY')
    run.font.size = Pt(14)
    run.bold = True
    heading_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.3)
    heading_cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.3)
    shading_grey = oxml.parse_xml(r'<w:shd {} w:fill="E1E1E1"/>'.
                                            format(oxml.ns.nsdecls('w')))
    heading_cells[0]._tc.get_or_add_tcPr().append(shading_grey)
    font.size = Pt(10)

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'FOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED.\n')
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.bold = True
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'UNVALIDATED FINDINGS MUST NOT BE ACTED UPON.\n')
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.bold = True
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'PLEASE CONTACT THE LABORATORY IF VALIDATION TESTING IS REQUIRED\n\n')
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.bold = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(
        0.3)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.3)
    
    run = table.rows[0].cells[0].paragraphs[0].add_run('Specialist Integrated Haematological Malignancy '
                                                       'Diagnostic Service\n')
    run.font.size = Pt(8)
    run = table.rows[0].cells[0].paragraphs[0].add_run('Acquired Genomics (SIHMDS-AG), Camelia Botnar Laboratories, '
                                                       'Great Ormond Street Hospital NHS Trust,\n')
    run.font.size = Pt(8)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'London, WC1N 3JH. Tel: 020 7405 9200 Ex: 5755')
    run.font.size = Pt(8)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(
        0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section. GENOMICS ENGLAND PARTICIPANT INFORMATION
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('GENOMICS ENGLAND PARTICIPANT INFORMATION')
    run.bold = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    shading_grey = oxml.parse_xml(r'<w:shd {} w:fill="E1E1E1"/>'.format(oxml.ns.nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_grey)

    table = document.add_table(rows=4, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'Patient Name:\t\t'
        f'{report.ir_family.participant_family.proband.forename} '
        f'{report.ir_family.participant_family.proband.surname}')
    run = table.rows[0].cells[1].paragraphs[0].add_run(
        f'NHS number:\t\t'
        f'{report.ir_family.participant_family.proband.nhs_number}')
    run = table.rows[1].cells[0].paragraphs[0].add_run(
        f'Date of Birth:\t\t'
        f'{report.ir_family.participant_family.proband.date_of_birth.date()}')
    run = table.rows[1].cells[1].paragraphs[0].add_run(
        f'CIP ID:\t\t\t'
        f'ILMN-{report.ir_family.ir_family_id}')
    run = table.rows[2].cells[0].paragraphs[0].add_run(
        f'Gender:\t\t'
        f'{report.ir_family.participant_family.proband.sex}')
    run = table.rows[2].cells[1].paragraphs[0].add_run(
        f'GEL Participant ID:\t'
        f'{report.ir_family.participant_family.proband.gel_id}')
    run = table.rows[3].cells[0].paragraphs[0].add_run(
        f'Referring Clinician:\t'
        f'{report.ir_family.participant_family.clinician.name}')
    run = table.rows[3].cells[1].paragraphs[0].add_run(
        f'Referring Hospital:\t'
        f'{report.ir_family.participant_family.proband.gmc}')
    
    # Section. CLINICAL UPDATE
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'CLINICAL UPDATE')
    run.bold = True
    shading_grey = oxml.parse_xml(r'<w:shd {} w:fill="E1E1E1"/>'.format(oxml.ns.nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_grey)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(
        0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'GTAB date: ')
    run = table.rows[0].cells[0].paragraphs[0].add_run('DD/MM/YYYY\n')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
    f'Information provided by: ')   
    run = table.rows[0].cells[0].paragraphs[0].add_run('XXXX (Title)\n')
    run.font.color.rgb = RGBColor(200, 200, 200)    
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'Include any SOC testing already performed, dates of '
        f'treatment, clinical trials etc.\n\n\n\n\n\n\n\n\n\n')
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    for cell in table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.space_before = Cm(0.2)
        cell.paragraphs[0].paragraph_format.space_after = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].add_run(
        f'Oncologist:\t\t')
    table.rows[0].cells[1].paragraphs[0].add_run(
        f'Clinician in lieu of referrer:\t\t')

    # Section. GENOMICS ENGLAND REPORT DETAILS
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'GENOMICS ENGLAND REPORT DETAILS')
    run.bold = True
    shading_grey = oxml.parse_xml(
        r'<w:shd {} w:fill="E1E1E1"/>'.format(oxml.ns.nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_grey)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    table = document.add_table(rows=7, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'Disease Type:\t\t\t'
        f'{report.ir_family.participant_family.proband.disease_group}')
    run = table.rows[0].cells[1].paragraphs[0].add_run(
        f'GEL Version Number:\t\t')
    run = table.rows[1].cells[0].paragraphs[0].add_run(
        f'Disease Subtype:\t\t'
        f'{report.ir_family.participant_family.proband.disease_subtype}')
    run = table.rows[1].cells[1].paragraphs[0].add_run(
        f'Multiple Samples:\t\t')
    run = table.rows[2].cells[0].paragraphs[0].add_run(
        f'Tumour Type:\t\t')
    run = table.rows[2].cells[1].paragraphs[0].add_run(
        f'% COSMIC Content < 30x:\t\t')
    run = table.rows[3].cells[0].paragraphs[0].add_run(
        f'Sample and Library Type:\t\t')
    run = table.rows[3].cells[1].paragraphs[0].add_run(
        f'Total Somatic SNVs:\t\t')
    run = table.rows[4].cells[0].paragraphs[0].add_run(
        f'Tumour Content:\t\t')      
    run = table.rows[4].cells[1].paragraphs[0].add_run(
        f'Total Somatic SVs:\t\t')
    run = table.rows[5].cells[0].paragraphs[0].add_run(
        f'Tumour contamination:\t\t ')
    run = table.rows[5].cells[1].paragraphs[0].add_run('1 ')
    run.font.superscript = True
    run = table.rows[5].cells[1].paragraphs[0].add_run(
        f'Mutation Burden:\t\t ')
    run = table.rows[6].cells[0].paragraphs[0].add_run(
        f'Date Analysis Issued:\t\t ')
    run = table.rows[6].cells[1].paragraphs[0].add_run('2 ')
    run.font.superscript = True
    run = table.rows[6].cells[1].paragraphs[0].add_run(
        f'Predominant Mutational Signature:\t\t')    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('1 ')
    run.font.superscript = True
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'Mutation Burden: >10 are classified as ‘hypermutators’ and >100, ‘ultra-hypermutators’.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run('2 ')
    run.font.superscript = True
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'Mutation Signature: is/is not notably found in disease type.\n'
    )
    
    # page break
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    # Page 2.
    # Section. GTAB SUMMARY
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'GTAB SUMMARY')
    run.bold = True
    run.font.size = Pt(10)
    shading_grey = oxml.parse_xml(
        r'<w:shd {} w:fill="E1E1E1"/>'.format(oxml.ns.nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_grey)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "\t\tFOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED\n\n")
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.bold = True
    run.font.size = Pt(10)

    run = table.rows[0].cells[0].paragraphs[0].add_run("SOMATIC VARIANTS\n")
    run.underline = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"Only variants with specific consequences (transcript ablation, splice acceptor variant, "
        f"splice donor variant, stop gained, frameshift variant, stop lost, start lost, transcript "
        f"amplification, inframe insertion, inframe deletion, inframe variant, missense variant, "
        f"splice region variant) in canonical transcripts are reported. Complex indels and frameshift "
        f"variants are only annotated at the coding sequence (CDS) level owing to problems accurately "
        f"annotating the protein change with the current pipeline. Small variants are classified as "
        f"single nucleotide variants (SNVs) and indels <50 base pairs (bp). Reported variants classified "
        f"into Domains 1-2 are reviewed by a clinical scientist and discussed at GTAB meeting.\n\n")
    run.font.size = Pt(8)
    
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"VAF, variant allele frequency. (LOH) a small variant that overlaps with copy-neutral loss of "
        f"heterozygosity region. Variant flag: (H) an indel intersecting with reference homopolymers of "
        f"at least 8 nucleotides in length (N) an indel in region with high levels of sequencing noise "
        f"where at least 10 % of the base calls in a window extending 50 bp either side of the indel's "
        f"call have been filtered out due to the poor quality (G) a variant with a germline allele "
        f"frequency > 1 % in an internal Genomics England data set(indicates potential unsubtracted "
        f"germline variant) (R) a recurrently identified somatic variant with somatic allele frequency "
        f"> 5 % in an internal Genomics England data set(indicates potential technical artefact) (SR) "
        f"a variant overlapping simple repeats. Variant flag(s) and/or low ALT allele frequency (<0.1), "
        f"variant is not scored.\n\n\n\n")
    run.font.size = Pt(8)

    run = table.rows[0].cells[0].paragraphs[0].add_run("Domain 0 ")
    run.underline = True
    run.font.bold = True
    run.font.size = Pt(10)

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"Please cut the variants and paste into their respective domains\n")
    run.underline = True
    run.font.size = Pt(10)

    count = 1
    for proband_variant in proband_variants:
        if proband_variant.max_tier == 0 and proband_variant.somatic is True:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n"
                                                         f"Genomic coordinate {proband_variant.variant.genome_assembly}"
                                                         f" ref>ALT allele: {proband_variant.variant.chromosome}:"
                                                         f"{proband_variant.variant.position} "
                                                         f"{proband_variant.variant.reference}>"
                                                         f"{proband_variant.variant.alternate}\n"
                                                         f"COSMIC ID and score:\n\n")
            count += 1

    run = table.rows[0].cells[0].paragraphs[0].add_run("Domain 1 ")
    run.underline = True
    run.font.bold = True
    run.font.size = Pt(10)

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"(variants in a panel of potentially actionable genes listed in the cancer "
        f"genome analysis documents, and reported as having therapeutic, prognostic or trial "
        f"associations by GenomOncology Knowledge Management System): \n\n\n\n")
    run.underline = True
    run.font.size = Pt(10)

    count = 1
    for proband_variant in proband_variants:
        if proband_variant.max_tier == 1 and proband_variant.somatic is True:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n"
                                                         f"Genomic coordinate {proband_variant.variant.genome_assembly}"
                                                         f"ref>ALT allele: {proband_variant.variant.chromosome}:"
                                                         f"{proband_variant.variant.position} "
                                                         f"{proband_variant.variant.reference}>"
                                                         f"{proband_variant.variant.alternate}\n"
                                                         f"COSMIC ID and score:\n\n")
            count += 1
    
    run = table.rows[0].cells[0].paragraphs[0].add_run("Domain 2 ")
    run.underline = True
    run.font.bold = True
    run.font.size = Pt(10)    
    
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"(variants in a broad panel of cancer-related genes listed in the cancer genome "
        f"analysis documents):\n\n")
    run.underline = True
    run.font.size = Pt(10)

    count = 1
    for proband_variant in proband_variants:
        if proband_variant.max_tier == 2 and proband_variant.somatic is True:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n"
                                                         f"Genomic coordinate {proband_variant.variant.genome_assembly}"
                                                         f"ref>ALT allele: {proband_variant.variant.chromosome}:"
                                                         f"{proband_variant.variant.position} "
                                                         f"{proband_variant.variant.reference}>"
                                                         f"{proband_variant.variant.alternate}\n"
                                                         f"COSMIC ID and score:\n\n")
            count += 1
    
    # Section. GERMLINE VARIANTS
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "\t\tFOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED\n\n")
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.font.bold = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"GERMLINE VARIANTS\n\n")
    run.underline = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"Germline variants in cancer susceptibility genes are divided into Tier 1 and Tier 3. "
        f"Tier 1 includes all truncating variants (in loss of function gene) and/or listed in "
        f"ClinVar to be pathogenic or likely pathogenic with a rating of at least two stars, detected "
        f"in cancer susceptibility genes relevant to the disease type. Tier 3 includes rare variants "
        f"that are not listed in ClinVar as benign or likely benign with a rating of at least two stars, "
        f"detected in a larger set of cancer susceptibility genes, listed in the cancer genome analysis "
        f"documents. Review of Tier 3 germline variants is not anticipated or required. Tier 3 variants "
        f"detected in cancer susceptibility genes relevant to the disease type are further investigated. "
        f"Interpretation of other Tier 3 variants can be requested if the variant is called in an "
        f"under-reported gene and/or in cases in which there is a high index of suspicion of a germline "
        f"determinant of cancer in the patient and/or family. Only germline variants classified as "
        f"‘likely pathogenic’ or ‘pathogenic’ are retained on the GTAB Summary Sheet at the point of "
        f"reporting.\n\n")
    run.font.size = Pt(8)

    run = table.rows[0].cells[0].paragraphs[0].add_run("Tier 1\n\n")
    run.font.underline = True
    run.font.bold = True
    run = table.rows[0].cells[0].paragraphs[0].add_run("Tier 3\n\n")
    run.font.underline = True
    run.font.bold = True
    count = 1
    for proband_variant in proband_variants:
        if proband_variant.somatic is False:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n"
                                                         f"Genomic coordinate {proband_variant.variant.genome_assembly}"
                                                         f"ref>ALT allele: {proband_variant.variant.chromosome}:"
                                                         f"{proband_variant.variant.position} "
                                                         f"{proband_variant.variant.reference}>"
                                                         f"{proband_variant.variant.alternate}\n"
                                                         f"COSMIC ID and score:\n\n")
            count += 1
    
    
    # Section. COPY NUMBER VARIANTS AND STRUCTURAL VARIANTS
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "\t\tFOR RESEARCH PURPOSES ONLY - THESE RESULTS HAVE NOT BEEN VALIDATED\n\n")
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.bold = True
    run.font.size = Pt(10)

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "COPY NUMBER VARIANTS AND STRUCTURAL VARIANTS\n\n")
    run.font.underline = True

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f"Somatic copy number variants(CNVs) and structural variants (SVs) that have not been assigned to "
        f"Domains (see ‘Somatic Variants’ ‘Domain 1’ and ‘Domain 2’ above) whilst the performance (recall "
        f"and precision) of the calling algorithm for CNVs and SVs is under evaluation. Only SVs overlapping "
        f"breakends with introns or exons are listed in the table below. Each row corresponds to one CNV or "
        f"SV. Types of variants: GAIN(COPY NUMBER) = CNV gain, LOSS(COPY NUMBER) = CNV loss, "
        f"LOH(COPY NUMBER)=loss of heterozygosity, BND=translocation, DEL=deletion, DUP=duplication, "
        f"INV=inversion and INS=insertion. Coordinate for the second breakend in translocation event "
        f"captures replacement string, position and direction according to variant call format specification v4.3.\n\n")
    run.font.size = Pt(8)
    
    run=table.rows[0].cells[0].paragraphs[0].add_run(
        "Translocations involving 2 or more named genes:\n\n\n\n")
    run.font.size = Pt(10)

    # Section. APPENDIX
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "\t\tFOR RESEARCH PURPOSES ONLY - THESE RESULTS HAVE NOT BEEN VALIDATED\n\n")
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    run.font.color.rgb = RGBColor(210, 42, 42)
    run.font.size = Pt(10)
    run.font.bold = True

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "APPENDIX\n\n")
    run.font.underline = True

    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'1. The analysis pipeline used filters out variants that fail quality '
        f'control thresholds, therefore absence of a variant does not indicate a '
        f'negative result.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'2. Technical information, including mutation burden calculation, for '
        f'the cancer genome analysis can be found at the Genomics England 100,000 '
        f'Genomes Project website - ')
    run_hyper = table.rows[0].cells[0].paragraphs[0].add_run(
        f'https://www.genomicsengland.co.uk/about-genomics-england/the-100000-genomes-project/information-for-gmc-staff/cancer-programme/cancer-genome-analysis\n')
    run_hyper.font.color.rgb = RGBColor(0, 0, 153)
    run_hyper.font.underline = True # has a hyperlink look, but not functional yet
    
    # TODO: write a method to add hyperlinks as part of the run?
    #run_hyper = table.rows[0].cells[0].paragraphs[0].add_run(
    #    f'https://www.genomicsengland.co.uk/about-genomics-england/the-100000-genomes-project/information-for-gmc-staff/cancer-programme/cancer-genome-analysis\n')
    #run_hyper.font.color.rgb = RGBColor(0, 0, 153)
    #run_hyper.font.underline = True
    #add_hyperlink_into_run(table.rows[0].cells[0].paragraphs[0], run_hyper, 'https://www.genomicsengland.co.uk/about-genomics-england/the-100000-genomes-project/information-for-gmc-staff/cancer-programme/cancer-genome-analysis')
   
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'3. Further details on mutational signatures, their prevalence in different '
        f'tumour types and proposed aetiology can be found at the Sanger COSMIC '
        f'signatures website - ')
    run_hyper = table.rows[0].cells[0].paragraphs[0].add_run(
        f'https://cancer.sanger.ac.uk/cosmic/signatures\n')
    run_hyper.font.color.rgb = RGBColor(0, 0, 153)
    run_hyper.font.underline = True # has a hyperlink look, but not functional yet   
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'4. Alamut Visual v2.11 is used for information on variant population '
        f'frequency, in silico prediction and splicing impact prediction.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'5. COSMIC IDs and gene-level actionability are sourced from the GEL '
        f'Associated Supplementary HTML.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'6. Somatic variant-level actionability is sourced from the GEL Associated '
        f'Supplementary HTML and/or if relevant to disease type \n')
    run_hyper = table.rows[0].cells[0].paragraphs[0].add_run(
        f'https://www.mycancergenome.org/.')
    run_hyper.font.color.rgb = RGBColor(0, 0, 153)
    run_hyper.font.underline = True # has a hyperlink look, but not functional yet
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f' Somatic Domain 1 categorisation is according to PMID: 27993330.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'7. Databases referenced include ClinVar, GeneCards, gnomAD, HGMD and OMIM.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'8. Germline predefined gene list v1.0 for each disease type can be provided upon request.\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'9. CNV/SV predefined gene list v1.0 for each disease type can be provided upon request.\n')

    # Section. CHECKED BY.
    table = document.add_table(rows=2, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('Completed by: ')
    run.bold = True
    run = table.rows[0].cells[0].paragraphs[0].add_run('XXXX (Title)')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run = table.rows[0].cells[1].paragraphs[0].add_run('Date: ')
    run.bold = True
    run = table.rows[0].cells[1].paragraphs[0].add_run('DD/MM/YYYY')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run = table.rows[1].cells[0].paragraphs[0].add_run('Checked by: ')
    run.bold = True
    run = table.rows[1].cells[0].paragraphs[0].add_run('XXXX (Title)')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run = table.rows[1].cells[1].paragraphs[0].add_run('Date: ')
    run.bold = True
    run = table.rows[1].cells[1].paragraphs[0].add_run('DD/MM/YYYY')
    run.font.color.rgb = RGBColor(200, 200, 200)

    return document


# TODO: Not in use currently. Planned for GTAB export append hyperlinks 
def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i+1,hyperlink)


def write_npf_template(report):
    '''
    Given a report, write a No Primary Findings (NPF) report
    :param report: GELInterpretation instance
    :return: docx document to be exported
    '''
    print("Report:", report)
    
    # template with headers, page number and custom Grid Table Plain setup
    template_file = os.path.join(os.getcwd(), "gel2mdt/exports_templates/{filename}".format(filename='npf_glh_negative_report_template.docx'))
    document = Document(docx=template_file)

    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # Demographics as a custom table style created in the docx template.
    # Setup gender and pronoun for text. If no gender, highlight in template
    try:
        # using gender pronoun in text
        sex = report.ir_family.participant_family.proband.sex
        clincian = report.ir_family.participant_family.clinician.name

        if not (sex is None or sex == 'unknown'):
            if sex.lower() == 'male':
                gender_pronoun = 'his'
                sex = list(report.ir_family.participant_family.proband.sex)[0].upper()
            elif sex.lower() == 'female':
                gender_pronoun = 'her'
                sex = list(report.ir_family.participant_family.proband.sex)[0].upper()
        else:
            sex = '<--Unknown-->'
            gender_pronoun = '<--his/her-->'
    except ValueError as e:
        print(e)
        raise

    # export letter date stamp
    now = datetime.now()

    # negative report function is avaliable for both sample types, this is
    # used within text for any future cancer requirement.
    if report.sample_type == 'raredisease':
        sample_type = 'rare disease'
    else:
        sample_type = 'cancer'
    
    table = document.add_table(rows=5, cols=2, style='Grid Table Plain')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        f'Dr {clincian}')
    run = table.rows[0].cells[1].paragraphs[0].add_run(
        f'Patient Name:\t\t  ')
    run.bold = True
    run = table.rows[0].cells[1].paragraphs[0].add_run(
        f'{report.ir_family.participant_family.proband.forename} '
        f'{report.ir_family.participant_family.proband.surname}')
    run = table.rows[1].cells[1].paragraphs[0].add_run(
        f'Date of Birth / Gender: ')
    run.bold = True
    run = table.rows[1].cells[1].paragraphs[0].add_run(
        f'{report.ir_family.participant_family.proband.date_of_birth.date().strftime("%d-%m-%Y")} / '
        f'{sex}')
    run = table.rows[2].cells[1].paragraphs[0].add_run(
        f'NHS number:\t\t  ')
    run.bold = True
    run = table.rows[2].cells[1].paragraphs[0].add_run(
        f'{report.ir_family.participant_family.proband.nhs_number}')
    run = table.rows[3].cells[1].paragraphs[0].add_run(
        f'GEL ID:\t\t  ')
    run.bold = True
    run = table.rows[3].cells[1].paragraphs[0].add_run(
        f'{report.ir_family.ir_family_id} / '
        f'{report.ir_family.participant_family.proband.gel_id}')
      
    # main text
    paragraph = document.add_paragraph()
    run = paragraph.add_run('\nNorth Thames Genomic Medicine Centre\n')
    run.font.size = Pt(14)
    run.bold = True
    run = paragraph.add_run('100,000 Genomes Project Result\n')
    run.font.size = Pt(11)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.bold = False

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'Date: '
        f'{now.strftime("%d/%m/%Y")}\n\n\n')
    
    # handle clincians entered without fullname
    if ' ' in clincian:
        clincian_surname = clincian.rsplit(' ', 1)[1]
    else:
        clincian_surname = clincian 
    
    run = paragraph.add_run(
        f'Dear Dr {clincian_surname},')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'The above named participant (and their family where applicable) are participating in the 100,000 Genomes  '
        f'Project to find the cause of {report.ir_family.participant_family.proband.forename}\'s {sample_type}. ')
    run = paragraph.add_run(
        f'Whole genome sequencing* has been completed by Genomics England and the primary analysis has not '
        f'identified any underlying genetic cause for {gender_pronoun} clinical presentation. ').bold = True
    run = paragraph.add_run(
        f'Please refer to the attached Genomics England report for information on the genes included in the primary '
        f'analysis. If panels appropriate to the phenotype have not been applied please contact the laboratory.')

    paragraph = document.add_paragraph(
        f'The genome sequencing data will be stored. Cases with new HPO terms, changing clinical need and those not '
        f'analysed for copy number variants (CNVs) will be re-analysed in the future as part of the on-going 100,000 '
        f'Genomes Project. If this identifies a possible genetic diagnosis we will re-contact you.')

    paragraph = document.add_paragraph(
        f'The analysis reported to date does not include analysis for ‘additional findings’ unrelated to the '
        f'clinical presentation. The results of the additional findings analysis will be reported separately '
        f'and at a later date to participants who have consented to receive these.')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'Please can you thank {report.ir_family.participant_family.proband.forename} (and their family where applicable) '
        f'for their continuing participation in the 100,000 Genomes Project. This letter should be stored in '
        f'{report.ir_family.participant_family.proband.forename}\'s medical records as a record of the result '
        f'and the outcome fed back to the participant(s) by the referring clinician.\n')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'Authorised by:\n\n\n')

    run = paragraph.add_run(
        f'Email: GEL.Team@gosh.nhs.uk\n')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f'* The whole genome sequencing analysis focussed on a panel of genes known to cause this patient’s condition '
        f'and is able to detect single nucleotide variants and small insertions/deletions. The analysis does not currently '
        f'detect larger copy number variants, deep intronic variants, structural abnormalities or variants on the '
        f'Y chromosome. Development and validation of software tools to identify such variants is in progress.')
    run.font.size = Pt(9)

    return document


def access_request_template():
    '''
    Another step added to the registration process.
    Formalising approval of user by line management.
    return: docx document to be exported
    '''
    template_filename = 'gel2mdt_access_request_template.docx'
    template_path = os.path.join(os.getcwd(), "gel2mdt/exports_templates/{filename}".format(filename=template_filename))
    document = Document(docx=template_path)
    return document